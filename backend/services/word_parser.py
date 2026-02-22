"""Parse contract Word (.docx) files to extract structured data."""

from io import BytesIO

from docx import Document

from backend.services.pdf_parser import _parse_basic_fields, _parse_fee_tables


def parse_contract_word(file_bytes: bytes) -> dict:
    result = {
        "contract_no": None,
        "company_name": None,
        "year": None,
        "start_date": None,
        "end_date": None,
        "total_amount": None,
        "raw_text": "",
        "water_plants": [],
        "ocr_used": False,
    }

    try:
        doc = Document(BytesIO(file_bytes))
    except Exception:
        result["raw_text"] = "[无法打开Word文件]"
        return result

    all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    all_tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            all_tables.append(rows)

    for table_rows in all_tables:
        for row in table_rows:
            row_text = " ".join(c for c in row if c)
            if row_text.strip():
                all_text += "\n" + row_text

    result["raw_text"] = all_text.strip()
    _parse_basic_fields(all_text, result)
    result["water_plants"] = _parse_fee_tables(all_tables)

    return result
